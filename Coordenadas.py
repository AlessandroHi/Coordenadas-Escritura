import os
import sys
import tkinter as tk
from tkinter import ttk
from customtkinter import *
from num2words import num2words
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from PIL import Image, ImageTk
import webbrowser
from num2words import num2words

def convertir_data(data):
    partes = data.split()
    
    punto_inicial = num2words(int(partes[0]), lang='es')
    punto_final = num2words(int(partes[2]), lang='es')
    
    azimut = partes[3].split('°')
    grados = num2words(int(azimut[0]), lang='es')
    
    minutos_raw = azimut[1].split("'")
    minutos = num2words(int(minutos_raw[0]), lang='es')
    
    segundos_raw = minutos_raw[1].strip('"').split('.')
    segundos = num2words(int(segundos_raw[0]), lang='es')
    
    distancia_entera = int(partes[4].split('.')[0])
    distancia_decimal = float('0.' + partes[4].split('.')[1])
    
    if distancia_entera == 1:
        distancia_str = "un metro"
    else:
        distancia_str = f"{num2words(distancia_entera, lang='es')} metros"
        
    distancia_decimal_str = num2words(int(distancia_decimal * 100), lang='es')
    
    azimut_str = f"azimut {grados} grados"
    if minutos != "cero":
        azimut_str += f", {minutos} minutos"
    if segundos != "cero":
        azimut_str += f", {segundos} segundos"
    
    if distancia_decimal != 0:
        distancia_str += f" y {distancia_decimal_str} centímetros exactos"
    
    resultado = f" Del punto {punto_inicial} al punto {punto_final}, {azimut_str} con distancia {distancia_str}."
    
    return resultado

def guardar_texto():
    parafo = ''
    # Obtener el texto del área de texto
    texto = textbox.get("1.0", "end-1c")  
    
    # Separar el texto en líneas
    lineas = texto.splitlines()
    
    # Recorrer cada línea
    for linea in lineas:
        # Convertir y procesar cada línea
        formato = convertir_data(linea)
        parafo += formato
   
    # Crear un nuevo documento de Word
    doc = Document()
    
    # Agregar un título
    doc.add_heading('', level=1)
    
    # Agregar texto
    paragraph = doc.add_paragraph(parafo)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    estilo_parrafo = doc.styles['Normal']
    estilo_parrafo.paragraph_format.line_spacing = Pt(20)
    
    # Establecer la fuente y el tamaño de la fuente
    estilo_parrafo.font.name = 'Calibri'
    estilo_parrafo.font.size = Pt(11)
    
    # Guardar el documento
    doc.save('Coordenadas.docx')

    # Construir la ruta completa del archivo
    try:
        os.startfile('Coordenadas.docx')  # Abre el archivo con la aplicación predeterminada
    except OSError as e:
        print(f"No se pudo abrir el documento: {e}")


def borrar_texto():
    textbox.delete("1.0", "end")


def open_github_profile():
    webbrowser.open("https://github.com/AlessandroHi")


    
if __name__ == "__main__":
    root = CTk()
    root.title("Tabla de Datos")
    root.geometry("700x600+300+100")    
    # Estilo de la ventana
    root.configure(bg="#A0A8AB")  # Cambia el color de fondo de la ventana

  
    textbox = CTkTextbox(master=root,height=500,width=400,corner_radius=15,scrollbar_button_color="#E8C25C",border_color="#EDE3CA",border_width=1,font=("Arial", 16),scrollbar_button_hover_color="#34C423")
    textbox.insert("0.0", "new text to insert")  # insert at line 0 character 0
    text = textbox.get("0.0", "end")  # get text from line 0 character 0 till the end
    textbox.delete("0.0", "end")  # delete all text
    textbox.configure(state="normal")  # configure textbox to be read-only
    textbox.place(x=220,y=20)

    # Crear el botón "Crear archivo" con los eventos
    agregar_button = CTkButton(master=root, text="Crear archivo",width=140,height=35, corner_radius=10,hover_color="#63C2E6",command=guardar_texto, fg_color="#219AC7", font=("Helvetica", 15))
    agregar_button.place(x=40, y=175)


    # Crear el botón "Borrar datos" con los eventos
    borrar_button = CTkButton(master=root, text="Borrar datos", width=140,height=35,corner_radius=10,hover_color="#EB5D4C",command=borrar_texto, fg_color="#A11807", font=("Helvetica", 15))
    borrar_button.place(x=42, y=235)

    footer_frame = tk.Frame(root, bg="#333333")
    footer_frame.pack(side="bottom", fill="x")

    footer_label = tk.Label(footer_frame, text="AlessandroHi", fg="white", bg="#333333", font=("Arial", 12))
    footer_label.pack(pady=5)
    
    my_image = CTkImage(light_image=Image.open("media/LOGO2.png"),size=(120, 120))
    image_label =CTkLabel(root, image=my_image, text="",fg_color="white",corner_radius=18)
    image_label.place(x=35,y=30)


    # Cargar la imagen del logo de GitHub
    github_logo = Image.open("media/imagen.png")
    github_logo = github_logo.resize((30, 30))
    github_logo = ImageTk.PhotoImage(github_logo)

    # Crear un Label para mostrar la imagen del logo de GitHub
    github_logo_label = tk.Label(footer_frame, image=github_logo, bg="#333333")
    github_logo_label.image = github_logo
    github_logo_label.pack(side="left", padx=5)

    # Centrar la imagen del logo de GitHub en el footer
    github_logo_label.place(relx=0.4, rely=0.5, anchor="center")
    # Enlazar la función al evento de clic en el nombre de usuario
    footer_label.bind("<Button-1>", lambda e: open_github_profile())

    # Enlazar la función al evento de clic en el logo de GitHub
    github_logo_label.bind("<Button-1>", lambda e: open_github_profile())


    root.mainloop()